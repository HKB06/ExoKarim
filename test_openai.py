import openai
import random
import datetime
import os
from docx import Document
from dotenv import load_dotenv

# 🔐 Chargement de la clé OpenAI depuis le fichier .env
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

# Charger les consignes depuis le document Word
chemin_fichier = "C:/Users/ps3ka/openai/BLOC1 DS.docx"
doc_consignes = Document(chemin_fichier)
texte_consignes = "\n".join([para.text for para in doc_consignes.paragraphs])

# Fonction pour introduire des erreurs subtiles
def introduire_erreurs_subtiles(texte):
    erreurs_subtiles = [
        ("inflation", "inflattion"),
        ("croissance économique", "croissanse economique"),
        ("régulation", "regulation"),
        ("volatilité", "volatilitée"),
        ("analyse financière", "analize financiere"),
        ("investissement", "investicement"),
        ("liquidité", "liquidé"),
        ("politique monétaire", "politique monnaitaire"),
        ("évaluation", "évalutation"),
        ("risque systémique", "riske systémique")
    ]
    if random.random() < 0.4:
        erreur = random.choice(erreurs_subtiles)
        texte = texte.replace(erreur[0], erreur[1], 1)
    return texte

# Génération des réponses via API OpenAI
def generer_reponses(questions, genre):
    responses = []
    for question in questions:
        temperature = random.uniform(0.7, 1.0)
        prompt = (
            f"Tu es un(e) étudiant(e) en finance. Réponds à cette question en utilisant un style personnel, comme si tu étais l'élève, à la première personne du singulier. "
            f"Utilise un style fluide, humain, avec des formulations naturelles. Tu es {'une étudiante' if genre == 'f' else 'un étudiant'}.\nQuestion : {question}"
        )
        response = openai.ChatCompletion.create(
            model="gpt-4",
            temperature=temperature,
            messages=[
                {"role": "system", "content": "Tu es un expert en finance de marché et veille stratégique."},
                {"role": "user", "content": prompt}
            ]
        )
        texte = response.choices[0].message.content
        responses.append(introduire_erreurs_subtiles(texte))
    return responses

# Informations élève
nom = input("Nom de l'élève : ")
prenom = input("Prénom de l'élève : ")
date_naissance = input("Date de naissance (JJ/MM/AAAA) : ")
genre = input("Genre de l'élève (m/f) : ").lower().strip()
date_realisation = datetime.datetime.today().strftime("%d/%m/%Y")

# Dossier de l'élève
repertoire_eleve = os.path.join("C:/Users/ps3ka/openai/Devoirs_Eleves", f"{prenom}_{nom}")
os.makedirs(repertoire_eleve, exist_ok=True)

# Extraction manuelle des parties et questions du document
devoirs = {
    "Partie 1": [
        "Réalisez une analyse SWOT pour évaluer l'impact des tensions commerciales sur le marché financier entre AATTLE et YIYO.",
        "Identifiez les nouvelles réglementations ou lois qui pourraient affecter les investisseurs en raison de ces tensions.",
        "Examinez les effets potentiels sur l'emploi, les inégalités sociales et autres aspects sociaux dans les deux pays et au niveau international.",
        "Analysez les implications environnementales de ces tensions, telles que les changements dans les chaînes d'approvisionnement qui pourraient affecter les émissions de carbone.",
        "Énoncez les outils de collecte et outils statistiques utilisés ainsi que la méthodologie pour l'analyse des indices boursiers."
    ],
        "Partie 2": [
        "Réalisez une étude de marché prospective (économique, juridique, sociale, environnementale) sur les tendances financières actuelles.",
        "Proposez une synthèse des principaux atouts et risques identifiés sur le marché.",
        "Analysez la volatilité récente des indices S&P 500, NASDAQ et CAC 40 à l’aide d’outils statistiques.",
        "Expliquez en quoi cette volatilité influence les décisions d’investissement.",
        "Présentez les outils statistiques utilisés pour cette analyse.",
        "Évaluez l’impact de la hausse des taux d’intérêt sur les actions, obligations et immobilier.",
        "Analysez l'effet de l’inflation sur divers secteurs et proposez des stratégies de couverture.",
        "Évaluez l’impact de l’inflation sur les secteurs technologique, énergétique et consommation de base et donnez une recommandation pour Apple dans chacun.",
        "Comparez les performances des secteurs : technologie, énergie, biens de consommation de base, santé, consommation discrétionnaire et immobilier.",
        "Proposez des hypothèses expliquant la sous-performance du secteur technologique."
    ],
    "Partie 3": [
        "Définissez une stratégie pour collecter des données sur les prix des actions des entreprises du CAC 40.",
        "Quels outils ou API allez-vous utiliser pour la collecte de ces données ?",
        "Proposez un algorithme de nettoyage et de préparation des données collectées.",
        "Quels indicateurs financiers allez-vous calculer (moyennes mobiles, RSI, etc.) ?",
        "Proposez un modèle statistique ou économétrique prédictif pour ces données.",
        "Évaluez la performance de votre modèle avec des indicateurs comme le R2 ajusté, le test de Student et le test de Fisher.",
        "Développez une interface utilisateur simple pour utiliser l’algorithme prédictif.",
        "Interprétez les résultats obtenus et préparez une présentation synthétique pour votre équipe."
    ],
    "Partie 4": [
        "Expliquez les avantages et inconvénients des modèles de valorisation suivants : Black-Scholes, modèle binomial, Monte-Carlo.",
        "Choisissez la meilleure méthode pour valoriser une option d’achat européenne (European Call) et justifiez votre choix.",
        "Calculez la valeur de l'option d’achat européenne avec la méthode choisie et expliquez toutes les étapes.",
        "Présentez différents scénarios de risque liés à cette option (volatilité, liquidité, taux d’intérêt, contrepartie, modèle).",
        "Estimez le temps nécessaire pour réaliser ces calculs avec votre méthode.",
        "Interprétez les résultats et expliquez leur implication pour la gestion des risques de la banque."
    ],
    "Partie 5": [
        "Analysez le profil d'investissement de Mme Dupont et justifiez pourquoi un produit structuré est adapté.",
        "Sélectionnez deux sous-jacents (actions, indices, matières premières) adaptés à son profil et aux conditions de marché.",
        "Concevez un produit structuré en décrivant la brique obligataire et la brique optionnelle, avec leurs caractéristiques (durée, protection, rendement).",
        "Présentez trois scénarios de rendement (favorable, neutre, défavorable) en chiffrant l’impact sur le capital de Mme Dupont.",
        "Vérifiez la conformité du produit structuré avec les réglementations (MiFID II, transparence, protection du capital).",
        "Rédigez une présentation pédagogique du produit structuré pour Mme Dupont, en expliquant clairement les avantages et les risques."
    ]

}

# Création du document
doc = Document()
doc.add_heading("Devoir Surveillé Complet - Analyse Financière et Stratégique", level=1)

# Informations élève
doc.add_paragraph(f"📌 Nom : {nom}")
doc.add_paragraph(f"📌 Prénom : {prenom}")
doc.add_paragraph(f"📌 Date de naissance : {date_naissance}")
doc.add_paragraph(f"📌 Date de réalisation : {date_realisation}")

# Traitement des parties
for partie, questions in devoirs.items():
    doc.add_heading(partie, level=1)
    reponses = generer_reponses(questions, genre)
    for idx, reponse in enumerate(reponses, 1):
        doc.add_heading(f"Question {idx}", level=2)
        doc.add_paragraph(reponse)

# Sauvegarde du fichier
filename = os.path.join(repertoire_eleve, f"Devoir_Surveillé_Complet_{prenom}_{nom}.docx")
doc.save(filename)

print(f"✅ Évaluation générée avec succès : '{filename}'")

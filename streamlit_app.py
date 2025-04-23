import streamlit as st
import openai
import random
import datetime
import os
import json
from docx import Document
from dotenv import load_dotenv

# 🔐 Chargement de la clé OpenAI depuis le fichier .env
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

# Charger les profils d’erreurs depuis le fichier JSON
with open("profils_fautes_corriges.json", "r", encoding="utf-8") as f:
    profils_fautes = json.load(f)

# 🔧 Génère un prompt plus naturel et académique
def construire_prompt(question, genre):
    intro = (
        "Tu incarnes un(e) étudiant(e) en Master 2 Finance de Marché – spécialité Trading. "
        "Tu rédiges une réponse à un devoir surveillé, en adoptant un ton personnel et critique, "
        "à la première personne du singulier, mais avec la rigueur attendue à ce niveau."
    )

    ton = (
        "Exprime-toi comme une étudiante en fin de Master, avec un style fluide, naturel et structuré. "
        "Utilise des transitions cohérentes, évite les répétitions comme 'En tant qu’étudiante…'. "
        "Sois précise, nuancée, et propose une réflexion argumentée, comme dans un devoir écrit authentique."
        if genre == "f" else
        "Exprime-toi comme un étudiant en fin de Master, avec un style fluide, naturel et structuré. "
        "Utilise des transitions cohérentes, évite les répétitions comme 'En tant qu’étudiant…'. "
        "Sois précis, nuancé, et propose une réflexion argumentée, comme dans un devoir écrit authentique."
    )

    consignes = (
        "Structure ta réponse avec : une courte introduction contextuelle, un développement logique, "
        "et une conclusion ou ouverture. N’utilise pas de listes à puces, mais développe chaque idée dans un paragraphe complet."
    )

    return f"{intro}\n{ton}\n{consignes}\n\nQuestion : {question}"

# 📌 Fonction pour introduire des fautes subtiles selon le profil
def introduire_erreurs_subtiles(texte, profil):
    base = profil["base_rate"]
    mots_risque = profil["mots_a_risque"]
    repetition = profil["repetition_faute"]

    if random.random() < base["accent_oublie"]:
        texte = texte.replace("é", "e", 1).replace("à", "a", 1)

    if random.random() < base["homophone"]:
        homophones = [("et", "est"), ("a", "à"), ("ces", "ses"), ("on", "ont"), ("c’est", "s’est")]
        h = random.choice(homophones)
        texte = texte.replace(h[0], h[1], 1)

    if random.random() < base["touche_adj"]:
        texte = texte.replace("u", "i", 1)

    if random.random() < base["liaison_oubliee"]:
        texte = texte.replace("n’est", "ne est", 1)

    if random.random() < base["fautes_d_accord"]:
        texte = texte.replace("les entreprises sont", "les entreprise est", 1)

    if random.random() < repetition:
        mot = random.choice(mots_risque)
        texte = texte.replace(mot, mot.upper(), 1)

    return texte

# 📌 Fonction pour générer les réponses
def generer_reponses(questions, genre, profil):
    responses = []
    for question in questions:
        temperature = random.uniform(0.7, 0.9)
        request_timeout=180,
        prompt = construire_prompt(question, genre)
        response = openai.ChatCompletion.create(
            model="gpt-4",
            temperature=temperature,
            messages=[
                {"role": "system", "content": "Tu es professeur et expert en finance de marché. Tu évalues un devoir rédigé par un(e) étudiant(e) en M2."},
                {"role": "user", "content": prompt}
            ]
        )
        texte = response.choices[0].message.content.strip()
        responses.append(introduire_erreurs_subtiles(texte, profil))
    return responses

# 🌐 Interface utilisateur
st.title("🧠 Générateur de Devoirs - Bloc 1")
st.markdown("Remplis le formulaire ci-dessous pour générer un devoir complet automatiquement.")

# 📋 Formulaire
nom = st.text_input("Nom de l'élève")
prenom = st.text_input("Prénom de l'élève")
date_naissance = st.date_input("Date de naissance")
genre = st.radio("Genre", ["m", "f"])

noms_profils = [profil["nom"] for profil in profils_fautes.values()]
profil_nom_selectionne = st.selectbox("Choisis un profil d'élève :", noms_profils)
profil_selectionne = next(p for p in profils_fautes.values() if p["nom"] == profil_nom_selectionne)


if st.button("📄 Générer le devoir"):
    date_realisation = datetime.datetime.today().strftime("%d/%m/%Y")
    repertoire_eleve = os.path.join("C:/Users/ps3ka/openai/Devoirs_Eleves", f"{prenom}_{nom}")
    os.makedirs(repertoire_eleve, exist_ok=True)

    devoirs = {
        "Partie 1": [
            "Réalisez une analyse SWOT pour évaluer l'impact des tensions commerciales sur le marché financier entre AATTLE et YIYO.",
            "Identifiez les nouvelles réglementations ou lois qui pourraient affecter les investisseurs en raison de ces tensions.",
            "Examinez les effets potentiels sur l'emploi, les inégalités sociales et autres aspects sociaux dans les deux pays et au niveau international.",
            "Analysez les implications environnementales de ces tensions, telles que les changements dans les chaînes d'approvisionnement qui pourraient affecter les émissions de carbone.",
            "Énoncez les outils de collecte et outils statistiques utilisés ainsi que la méthodologie pour l'analyse des indices boursiers."
        ],
        "Partie 2 – Étude des marchés financiers et stratégies d’investissement": [
            "Réalisez une étude de marché prospective (économique, juridique, sociale, environnementale) sur les tendances financières actuelles.",
            "Proposez une synthèse des principaux atouts et risques identifiés sur le marché.",
            "Analysez la volatilité récente des indices S&P 500, NASDAQ et CAC 40 à l’aide d’outils statistiques.",
            "Expliquez en quoi cette volatilité influence les décisions d’investissement.",
            "Présentez les outils statistiques utilisés pour cette analyse.",
            "Évaluez l’impact de la hausse des taux d’intérêt sur les actions, obligations et immobilier.",
            "Analysez l'effet de l’inflation sur divers secteurs et proposez des stratégies de couverture.",
            "Évaluez l’impact de l’inflation sur les secteurs technologique, énergétique et consommation de base et donnez une recommandation pour Apple dans chacun.",
            "Comparez les performances des secteurs : technologie, énergie, biens de consommation de base, santé, consommation discrétionnaire et immobilier.",
            "Proposez des hypothèses expliquant la sous-performance du secteur technologique."
        ],
        "Partie 3 – Stratégie de collecte et analyse de données": [
            "Définissez une stratégie pour collecter des données sur les prix des actions des entreprises du CAC 40.",
            "Quels outils ou API allez-vous utiliser pour la collecte de ces données ?",
            "Proposez un algorithme de nettoyage et de préparation des données collectées.",
            "Quels indicateurs financiers allez-vous calculer (moyennes mobiles, RSI, etc.) ?",
            "Proposez un modèle statistique ou économétrique prédictif pour ces données.",
            "Évaluez la performance de votre modèle avec des indicateurs comme le R2 ajusté, le test de Student et le test de Fisher.",
            "Développez une interface utilisateur simple pour utiliser l’algorithme prédictif.",
            "Interprétez les résultats obtenus et préparez une présentation synthétique pour votre équipe."
        ],
        "Partie 4 – Valorisation des instruments financiers": [
            "Expliquez les avantages et inconvénients des modèles de valorisation suivants : Black-Scholes, modèle binomial, Monte-Carlo.",
            "Choisissez la meilleure méthode pour valoriser une option d’achat européenne (European Call) et justifiez votre choix.",
            "Calculez la valeur de l'option d’achat européenne avec la méthode choisie et expliquez toutes les étapes.",
            "Présentez différents scénarios de risque liés à cette option (volatilité, liquidité, taux d’intérêt, contrepartie, modèle).",
            "Estimez le temps nécessaire pour réaliser ces calculs avec votre méthode.",
            "Interprétez les résultats et expliquez leur implication pour la gestion des risques de la banque."
        ],
        "Partie 5 – Conseil en investissement et produits structurés": [
            "Analysez le profil d'investissement de Mme Dupont et justifiez pourquoi un produit structuré est adapté.",
            "Sélectionnez deux sous-jacents (actions, indices, matières premières) adaptés à son profil et aux conditions de marché.",
            "Concevez un produit structuré en décrivant la brique obligataire et la brique optionnelle, avec leurs caractéristiques (durée, protection, rendement).",
            "Présentez trois scénarios de rendement (favorable, neutre, défavorable) en chiffrant l’impact sur le capital de Mme Dupont.",
            "Vérifiez la conformité du produit structuré avec les réglementations (MiFID II, transparence, protection du capital).",
            "Rédigez une présentation pédagogique du produit structuré pour Mme Dupont, en expliquant clairement les avantages et les risques."
        ]
    }

    doc = Document()
    doc.add_heading("Devoir Surveillé Complet - Analyse Financière et Stratégique", level=1)
    doc.add_paragraph(f"📌 Nom : {nom}")
    doc.add_paragraph(f"📌 Prénom : {prenom}")
    doc.add_paragraph(f"📌 Date de naissance : {date_naissance.strftime('%d/%m/%Y')}")
    doc.add_paragraph(f"📌 Date de réalisation : {date_realisation}")

    for partie, questions in devoirs.items():
        doc.add_heading(partie, level=1)
        reponses = generer_reponses(questions, genre,profil_selectionne)
        for idx, reponse in enumerate(reponses, 1):
            doc.add_heading(f"Question {idx}", level=2)
            doc.add_paragraph(reponse)

    filename = os.path.join(repertoire_eleve, f"Devoir_Surveillé_Complet_{prenom}_{nom}.docx")
    doc.save(filename)

    st.success(f"✅ Devoir généré et enregistré : {filename}")
    with open(filename, "rb") as file:
        st.download_button("⬇️ Télécharger le devoir", file, file_name=os.path.basename(filename))

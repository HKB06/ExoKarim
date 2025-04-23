import streamlit as st
import openai
import random
import datetime
import os
import json
from docx import Document

# Configuration de la clé OpenAI
openai.api_key = os.getenv("OPENAI_API_KEY")

# Chargement des profils d’erreurs
with open("profils_fautes_corriges.json", "r", encoding="utf-8") as f:
    profils_fautes = json.load(f)

# Fonction pour construire un prompt adapté

def construire_prompt(question, genre):
    intro = (
        "Tu incarnes un(e) étudiant(e) en Master 2 Finance de Marché – spécialité Trading. "
        "Tu rédiges une réponse à un devoir surveillé, en adoptant un ton personnel et critique, "
        "à la première personne du singulier, mais avec la rigueur attendue à ce niveau."
    )

    ton = (
        "Exprime-toi comme une étudiante en fin de Master, avec un style fluide, naturel et structuré. "
        "Utilise des transitions cohérentes, évite les répétitions comme 'En tant qu’étudiante…'. "
        "Sois précise, nuancée, et propose une réflexion argumentée, comme dans un devoir écrit authentique."
        if genre == "f" else
        "Exprime-toi comme un étudiant en fin de Master, avec un style fluide, naturel et structuré. "
        "Utilise des transitions cohérentes, évite les répétitions comme 'En tant qu’étudiant…'. "
        "Sois précis, nuancé, et propose une réflexion argumentée, comme dans un devoir écrit authentique."
    )

    consignes = (
        "Structure ta réponse avec : une courte introduction contextuelle, un développement logique, "
        "et une conclusion ou ouverture. N’utilise pas de listes à puces, mais développe chaque idée dans un paragraphe complet."
    )

    return f"{intro}\n{ton}\n{consignes}\n\nQuestion : {question}"

# Fonction pour introduire des erreurs selon un profil

def introduire_erreurs_subtiles(texte, profil):
    base = profil["base_rate"]
    mots_risque = profil["mots_a_risque"]
    repetition = profil["repetition_faute"]

    if random.random() < base["accent_oublie"]:
        texte = texte.replace("é", "e", 1).replace("à", "a", 1)

    if random.random() < base["homophone"]:
        homophones = [("et", "est"), ("a", "à"), ("ces", "ses"), ("on", "ont"), ("c’est", "s’est")]
        h = random.choice(homophones)
        texte = texte.replace(h[0], h[1], 1)

    if random.random() < base["touche_adj"]:
        texte = texte.replace("u", "i", 1)

    if random.random() < base["liaison_oubliee"]:
        texte = texte.replace("n’est", "ne est", 1)

    if random.random() < base["fautes_d_accord"]:
        texte = texte.replace("les entreprises sont", "les entreprise est", 1)

    if random.random() < repetition:
        mot = random.choice(mots_risque)
        texte = texte.replace(mot, mot.upper(), 1)

    return texte

# Chargement du fichier devoirs_bloc2.json avec contexte, tâches et questions
with open("devoirs_bloc2.json", "r", encoding="utf-8") as f:
    devoirs_bloc2 = json.load(f)

# Génération des réponses par devoir

def generer_reponses_par_bloc(partie_info, genre, profil):
    contexte = partie_info["contexte"]
    taches = partie_info["taches"]
    questions = partie_info["questions"]
    responses = []

    for question in questions:
        prompt = (
            f"Contexte :\n{contexte.strip()}\n\n"
            f"Tâche à accomplir :\n{taches.strip()}\n\n"
            f"Question :\n{question.strip()}\n\n"
            f"{construire_prompt(question, genre)}"
        )

        response = openai.ChatCompletion.create(
            model="gpt-4",
            temperature=random.uniform(0.7, 0.9),
            messages=[
                {"role": "system", "content": "Tu es professeur en finance de marché, expert en stratégie macro-économique et évaluation financière."},
                {"role": "user", "content": prompt}
            ]
        )

        texte = response.choices[0].message.content.strip()
        texte_avec_fautes = introduire_erreurs_subtiles(texte, profil)
        responses.append(texte_avec_fautes)

    return responses

# Interface utilisateur Streamlit

st.title("🧠 Générateur de Devoirs - Bloc 2")
st.markdown("Remplis le formulaire ci-dessous pour générer un devoir complet automatiquement.")

nom = st.text_input("Nom de l'élève")
prenom = st.text_input("Prénom de l'élève")
date_naissance = st.date_input("Date de naissance")
genre = st.radio("Genre", ["m", "f"])

noms_profils = [profil["nom"] for profil in profils_fautes.values()]
profil_nom_selectionne = st.selectbox("Choisis un profil d'élève :", noms_profils)
profil_selectionne = next(p for p in profils_fautes.values() if p["nom"] == profil_nom_selectionne)

if st.button("📄 Générer le devoir"):
    date_realisation = datetime.datetime.today().strftime("%d/%m/%Y")
    repertoire_eleve = os.path.join("Devoirs_Eleves_Bloc2", f"{prenom}_{nom}")
    os.makedirs(repertoire_eleve, exist_ok=True)

    doc = Document()
    doc.add_heading("Devoir Surveillé Complet - Bloc 2 - Master Finance de Marché", level=1)
    doc.add_paragraph(f"📌 Nom : {nom}")
    doc.add_paragraph(f"📌 Prénom : {prenom}")
    doc.add_paragraph(f"📌 Date de naissance : {date_naissance.strftime('%d/%m/%Y')}")
    doc.add_paragraph(f"📌 Date de réalisation : {date_realisation}")

    for partie, contenu in devoirs_bloc2.items():
        doc.add_page_break()
        doc.add_heading(f"📘 {partie}", level=1)
        reponses = generer_reponses_par_bloc(contenu, genre, profil_selectionne)
        for idx, reponse in enumerate(reponses, 1):
            doc.add_heading(f"📝 Question {idx}", level=2)
            doc.add_paragraph(reponse)

    filename = os.path.join(repertoire_eleve, f"Devoir_Surveillé_Complet_{prenom}_{nom}.docx")
    doc.save(filename)

    st.success(f"✅ Devoir généré et enregistré : {filename}")
    with open(filename, "rb") as file:
        st.download_button("⬇️ Télécharger le devoir", file, file_name=os.path.basename(filename))
